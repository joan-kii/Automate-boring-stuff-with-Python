#!python3
#customInvitations. Crea invitaciones para una lista de invitados.

import docx

# Abre el archivo 'guests.txt'.
guestsFile = open('guests.txt', 'r')
# Convierte el texto en listas.
invitados = guestsFile.readlines()
# Carga el documento '.docx' con los estilos escogidos.
doc = docx.Document('C:\\Users\\tu_path\\guests.docx')

# Itera sobre la lista 'invitados' y escribe los distintos párrafos con
# los distintos estilos creados previamente en el documento 
# 'guests.docx'. Crea un archivo con cada invitación.
for invitado in invitados:
    # Elimina los espacios del nombre de invitado.
    invitado.strip()

    p1 = doc.add_paragraph('Sería un placer para nosotros poder disfrutar de la compañía de', 'Estilo1')
    p2 = doc.add_paragraph(invitado, 'Estilo2')
    p3 = doc.add_paragraph('en el Castillo Milenario, Narnia', 'Estilo1')
    p4 = doc.add_paragraph('el día 20 de Abril del 2020', 'Estilo1')
    p5 = doc.add_paragraph('a las 20.00', 'Estilo3')

    # Salta a otra página para crear una invitación por página.
    doc.add_page_break()

# Guarda el documento.
doc.save('customInvitations.docx')